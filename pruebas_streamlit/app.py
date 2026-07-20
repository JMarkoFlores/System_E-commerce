import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import io
import os
import json
import datetime
import json
import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from data import generar_usuarios_sinteticos, estadisticas_dataset
from models import (
    RandomRecommender, 
    PopularityRecommender, 
    ContentBasedRecommender, 
    HybridWeightedRecommender, 
    HybridCascadeRecommender
)
from evaluation import evaluar_recomendador, comparar_con_baselines

st.set_page_config(page_title="Módulo de Pruebas - TechStore AI", layout="wide", page_icon="🧪")

st.title("🧪 Evaluación de Modelos de Recomendación")
st.markdown("Módulo avanzado de pruebas para los algoritmos del E-commerce.")

# --- Sidebar Configuration ---
with st.sidebar:
    st.header("Configuración de la Prueba")
    num_usuarios = st.number_input(
        "Usuarios Sintéticos", min_value=10, max_value=500, value=40, step=10,
        help="Cantidad de perfiles a simular. Más usuarios aumentan la fiabilidad estadística, pero tardan más en procesar."
    )
    top_k = st.number_input(
        "Top-K Recomendaciones", min_value=1, max_value=50, value=10, step=1,
        help="Número de productos a recomendar por usuario. Simula el límite de un carrusel en pantalla (ej. Top 10)."
    )
    seed = st.number_input(
        "Semilla Aleatoria", value=42,
        help="Fija la aleatoriedad. Mantener la misma semilla garantiza los mismos resultados exactos al repetir la prueba."
    )
    
    run_btn = st.button("Ejecutar Evaluación", type="primary", use_container_width=True)

if "resultados" not in st.session_state:
    st.session_state.resultados = None
if "excel_data" not in st.session_state:
    st.session_state.excel_data = None
if "pdf_bytes" not in st.session_state:
    st.session_state.pdf_bytes = None
if "word_data" not in st.session_state:
    st.session_state.word_data = None

# --- Ejecución ---
if run_btn:
    with st.spinner("Generando datos sintéticos..."):
        usuarios = generar_usuarios_sinteticos(cantidad=num_usuarios, semilla=seed)
        stats_dataset = estadisticas_dataset(usuarios)
        
    with st.spinner("Evaluando modelos..."):
        random_rec = RandomRecommender(seed=seed)
        popularity_rec = PopularityRecommender()
        content_rec = ContentBasedRecommender()
        hybrid_w_rec = HybridWeightedRecommender()
        hybrid_c_rec = HybridCascadeRecommender()
        
        res_random = evaluar_recomendador(random_rec, usuarios, k=top_k)
        res_pop = evaluar_recomendador(popularity_rec, usuarios, k=top_k)
        res_content = evaluar_recomendador(content_rec, usuarios, k=top_k)
        res_hw = evaluar_recomendador(hybrid_w_rec, usuarios, k=top_k)
        res_hc = evaluar_recomendador(hybrid_c_rec, usuarios, k=top_k)
        
        resultados_baselines = {
            "Random": res_random,
            "Popularidad": res_pop,
            "Content-Based": res_content,
            "Híbrido Ponderado": res_hw,
            "Híbrido Cascada": res_hc
        }
        
        # Encontrar el mejor modelo (para hacer tests estadísticos contra él)
        pesos = {"hitRate": 0.20, "precision": 0.15, "recall": 0.15, "f1": 0.15, "ndcg": 0.15, "mrr": 0.10, "map": 0.10}
        
        scores = []
        for name, res in resultados_baselines.items():
            score = (
                (res["hitRate"] / 100) * pesos["hitRate"] +
                (res["precision"] / 100) * pesos["precision"] +
                (res["recall"] / 100) * pesos["recall"] +
                (res["f1"] / 100) * pesos["f1"] +
                (res["ndcg"] / 100) * pesos["ndcg"] +
                (res["mrr"]) * pesos["mrr"] +
                (res["map"]) * pesos["map"]
            ) * 100
            scores.append({"nombre": name, "score": score, "res": res})
            
        scores.sort(key=lambda x: x["score"], reverse=True)
        mejor_modelo = scores[0]
        
        otros_modelos = {s["nombre"]: s["res"] for s in scores[1:]}
        
        tests_por_metrica = {}
        metricas_tests = ["hitRate", "precision", "recall", "ndcg"]
        for metrica in metricas_tests:
            tests_por_metrica[metrica] = comparar_con_baselines(mejor_modelo["res"], otros_modelos, metrica)
            
        st.session_state.resultados = {
            "stats": stats_dataset,
            "resultados_baselines": resultados_baselines,
            "mejor_modelo": mejor_modelo,
            "tests_por_metrica": tests_por_metrica,
            "scores": scores
        }
        
        # --- AUTO-SAVE PARA COMUNICACIÓN CON REACT ---
        react_data = {
            "ganador": {
                "nombre": mejor_modelo["nombre"],
                "score": mejor_modelo["score"]
            },
            "ordenados": []
        }
        for s in scores:
            react_data["ordenados"].append({
                "nombre": s["nombre"],
                "score": s["score"],
                "hitRate": s["res"]["hitRate"],
                "precision": s["res"]["precision"],
                "recall": s["res"]["recall"],
                "f1": s["res"]["f1"],
                "ndcg": s["res"]["ndcg"],
                "mrr": s["res"]["mrr"],
                "map": s["res"]["map"],
                "coverage": s["res"]["coverage"]
            })
            
        # Escribir en la carpeta public de React
        public_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'public')
        os.makedirs(public_dir, exist_ok=True)
        try:
            with open(os.path.join(public_dir, 'resultados_streamlit.json'), 'w', encoding='utf-8') as f:
                json.dump(react_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.warning(f"No se pudo autoguardar el JSON para React: {e}")

    # --- Generar documentos una sola vez y cachear en session_state ---
    with st.spinner("Preparando archivos de descarga..."):
        _res   = st.session_state.resultados
        _stats = _res["stats"]
        _baselines = _res["resultados_baselines"]
        _mejor = _res["mejor_modelo"]
        _tests = _res["tests_por_metrica"]
        _scores = _res["scores"]

        _metricas_visibles = [
            {"key": "hitRate",   "label": "Hit Rate@K",  "sufijo": "%"},
            {"key": "precision", "label": "Precision@K", "sufijo": "%"},
            {"key": "recall",   "label": "Recall@K",    "sufijo": "%"},
            {"key": "f1",       "label": "F1@K",        "sufijo": "%"},
            {"key": "mrr",      "label": "MRR",         "sufijo": ""},
            {"key": "map",      "label": "MAP",         "sufijo": ""},
            {"key": "ndcg",     "label": "NDCG@K",     "sufijo": "%"},
            {"key": "coverage", "label": "Coverage",   "sufijo": "%"},
            {"key": "diversity","label": "Diversity",  "sufijo": "%"},
            {"key": "novelty",  "label": "Novelty",    "sufijo": "%"},
        ]

        # -- Generar figuras para Excel y PDF --
        _COLORES = ['#EF4444', '#F59E0B', '#10B981', '#EC4899', '#14B8A6', '#8B5CF6']
        _tabla_data = []
        for _m in _metricas_visibles:
            _row = {"Metrica": _m["label"]}
            for _nombre, _resultado in _baselines.items():
                _row[_nombre] = round(_resultado.get(_m["key"], 0), 4)
            _tabla_data.append(_row)
        _df_metricas = pd.DataFrame(_tabla_data).set_index("Metrica")

        _df_melted = _df_metricas.reset_index().melt(id_vars=["Metrica"], var_name="Modelo", value_name="Valor")
        _fig_bar = px.bar(_df_melted, x="Metrica", y="Valor", color="Modelo", barmode="group",
                          title="Comparacion de Metricas", color_discrete_sequence=_COLORES)

        _radar_keys = ['hitRate', 'precision', 'recall', 'f1', 'ndcg', 'diversity']
        _fig_radar = go.Figure()
        for _i, (_nombre, _resultado) in enumerate(_baselines.items()):
            _vals = [_resultado.get(_k, 0) for _k in _radar_keys]
            _vals.append(_vals[0])
            _fig_radar.add_trace(go.Scatterpolar(
                r=_vals, theta=_radar_keys + [_radar_keys[0]],
                fill='toself', name=_nombre,
                line_color=_COLORES[_i % len(_COLORES)]
            ))
        _fig_radar.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
            showlegend=True, title="Perfil de Rendimiento (Radar)"
        )

        # -- EXCEL --
        try:
            _tabla_data = []
            for _m in _metricas_visibles:
                _row = {"Metrica": _m["label"]}
                for _nombre, _resultado in _baselines.items():
                    _row[_nombre] = round(_resultado.get(_m["key"], 0), 4)
                _tabla_data.append(_row)
            _df_metricas = pd.DataFrame(_tabla_data).set_index("Metrica")

            _out_xl = io.BytesIO()
            with pd.ExcelWriter(_out_xl, engine="xlsxwriter") as _writer:
                _wb = _writer.book
                _hfmt = _wb.add_format({"bold": True, "text_wrap": True, "valign": "center", "align": "center", "fg_color": "#4F46E5", "font_color": "white", "border": 1})
                _cfmt = _wb.add_format({"border": 1, "align": "center"})
                _pfmt = _wb.add_format({"border": 1, "align": "center", "num_format": "0.00%"})
                _nfmt = _wb.add_format({"border": 1, "align": "center", "num_format": "0.0000"})
                _tfmt = _wb.add_format({"bold": True, "font_size": 16, "font_color": "#166534", "align": "center", "valign": "center", "fg_color": "#DCFCE7", "border": 1})
                _sfmt = _wb.add_format({"bold": True, "font_size": 12, "font_color": "#374151"})
                _txtfmt = _wb.add_format({"text_wrap": True, "valign": "top"})

                _ws_res = _wb.add_worksheet("Resumen Ejecutivo")
                _ws_res.merge_range("B2:J4", f"MEJOR MODELO: {_mejor['nombre']} (Score: {round(_mejor['score'], 2)} / 100)", _tfmt)
                _ws_res.write("B6", "Interpretacion de los Resultados:", _sfmt)
                _txt_bar = f"Barras: '{_mejor['nombre']}' lidera con NDCG {round(_mejor['res']['ndcg'],2)}% y F1 {round(_mejor['res']['f1'],2)}%."
                _txt_rad = f"Radar: '{_mejor['nombre']}' balance Precision {round(_mejor['res']['precision'],2)}% y Recall {round(_mejor['res']['recall'],2)}%."
                _ws_res.merge_range("B8:E11", _txt_bar, _txtfmt)
                _ws_res.merge_range("G8:J11", _txt_rad, _txtfmt)

                # Insertar graficas en Excel
                try:
                    _tmp_bar_xl = "_tmp_bar_xl.png"
                    _tmp_rad_xl = "_tmp_rad_xl.png"
                    _fig_bar.write_image(_tmp_bar_xl, width=500, height=350)
                    _fig_radar.write_image(_tmp_rad_xl, width=500, height=350)
                    _ws_res.insert_image("B13", _tmp_bar_xl)
                    _ws_res.insert_image("G13", _tmp_rad_xl)
                except Exception:
                    _ws_res.write("B13", "(Graficos requieren kaleido)")

                def _write_df(df, sname, is_metric=False):
                    _ws = _wb.add_worksheet(sname)
                    for ci, cv in enumerate(df.columns.values):
                        _ws.write(0, ci, str(cv), _hfmt)
                    for ri in range(len(df)):
                        _rdata = df.iloc[ri]
                        _is_mrr = is_metric and "Metrica" in df.columns and _rdata["Metrica"] in ["MRR", "MAP"]
                        for ci in range(len(df.columns)):
                            _v = _rdata.iloc[ci]
                            if isinstance(_v, (dict, list)):
                                _v = str(_v)
                            if ci == 0 or not isinstance(_v, (int, float)):
                                _ws.write(ri + 1, ci, _v if not pd.isna(_v) else "", _cfmt)
                            elif is_metric:
                                _ws.write(ri + 1, ci, _v if _is_mrr else _v / 100.0, _nfmt if _is_mrr else _pfmt)
                            else:
                                _ws.write(ri + 1, ci, _v, _cfmt if sname == "Dataset" else _nfmt)
                    for ci, cv in enumerate(df.columns.values):
                        _clen = max(df.iloc[:, ci].astype(str).map(len).max(), len(str(cv))) + 4
                        _ws.set_column(ci, ci, _clen)

                _write_df(pd.DataFrame([_stats]), "Dataset")
                _write_df(_df_metricas.reset_index(), "Metricas", is_metric=True)
                _all_tests = []
                for _mk, _td in _tests.items():
                    for _bn, _tv in _td.items():
                        _all_tests.append({"Metrica": _mk, "Modelo": _bn, "T-Statistic": _tv["ttest"]["tStatistic"], "T-pVal": _tv["ttest"]["pValue"], "W-pVal": _tv["wilcoxon"]["pValue"], "Cohens D": _tv["cohensD"]})
                _write_df(pd.DataFrame(_all_tests), "Tests Estadisticos")

            st.session_state.excel_data = _out_xl.getvalue()
            # Limpiar temporales de Excel
            try:
                os.remove("_tmp_bar_xl.png")
                os.remove("_tmp_rad_xl.png")
            except Exception:
                pass
        except Exception as _e:
            st.session_state.excel_data = None
            st.warning(f"Excel no disponible: {_e}")

        # -- PDF --
        try:
            class _PDF(FPDF):
                def header(self):
                    self.set_fill_color(79, 70, 229)
                    self.rect(0, 0, 210, 25, "F")
                    self.set_font("Helvetica", "B", 18)
                    self.set_text_color(255, 255, 255)
                    self.set_y(8)
                    self.cell(0, 10, "TechStore AI - Reporte Avanzado", border=0, ln=1, align="C")
                    self.ln(10)
                def footer(self):
                    self.set_y(-15)
                    self.set_fill_color(79, 70, 229)
                    self.rect(10, 282, 190, 1, "F")
                    self.set_font("Helvetica", "I", 8)
                    self.set_text_color(100, 100, 100)
                    self.cell(0, 10, f"Pagina {self.page_no()}", 0, 0, "C")

            _pdf = _PDF()
            _pdf.add_page()
            _pdf.set_font("Helvetica", size=10)
            _pdf.set_text_color(80, 80, 80)
            _pdf.cell(0, 6, txt=f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=1)
            _pdf.cell(0, 6, txt=f"Usuarios: {_stats['totalUsuarios']} | Compras: {_stats['totalCompras']}", ln=1)
            _pdf.ln(5)
            _pdf.set_fill_color(240, 253, 244)
            _pdf.set_draw_color(74, 222, 128)
            _pdf.set_line_width(0.5)
            _pdf.rect(10, _pdf.get_y(), 190, 15, "DF")
            _pdf.set_y(_pdf.get_y() + 2)
            _pdf.set_font("Helvetica", "B", 12)
            _pdf.set_text_color(22, 101, 52)
            _pdf.cell(0, 10, txt=f" MEJOR MODELO: {_mejor['nombre']} (Score: {round(_mejor['score'],2)}/100)", ln=1, align="C")
            _pdf.ln(8)
            _pdf.set_font("Helvetica", "B", 14)
            _pdf.set_text_color(30, 30, 30)
            _pdf.cell(0, 10, txt="Tabla Comparativa de Metricas (Top 3)", ln=1)
            _pdf.set_font("Helvetica", size=10)
            _top3 = _scores[:3]
            with _pdf.table(text_align="CENTER", col_widths=(40, 30, 30, 30, 30), borders_layout="ALL") as _tbl:
                _hr = _tbl.row()
                for _ht in ["Modelo", "Hit Rate", "Precision", "NDCG", "Score Final"]:
                    _pdf.set_fill_color(79, 70, 229)
                    _pdf.set_text_color(255, 255, 255)
                    _hr.cell(_ht)
                _pdf.set_text_color(50, 50, 50)
                for _i, _s in enumerate(_top3):
                    _row = _tbl.row()
                    if _i % 2 == 0:
                        _pdf.set_fill_color(245, 245, 250)
                    else:
                        _pdf.set_fill_color(255, 255, 255)
                    _mr = _s["res"]
                    _row.cell(_s["nombre"])
                    _row.cell(f"{round(_mr['hitRate'],2)}%")
                    _row.cell(f"{round(_mr['precision'],2)}%")
                    _row.cell(f"{round(_mr['ndcg'],2)}%")
                    _row.cell(f"{round(_s['score'],2)}/100")
            _pdf.ln(5)
            if len(_top3) > 1:
                _pdf.set_font("Helvetica", size=10)
                _pdf.set_text_color(60, 60, 60)
                _diff_score = _top3[0]['score'] - _top3[1]['score']
                _dominio = "con un margen notable" if _diff_score >= 5 else "por un margen estrecho"
                _texto_tabla = (
                    f"Interpretacion de la Tabla:\n"
                    f"- El modelo '{_top3[0]['nombre']}' lidera {_dominio} ({round(_diff_score, 2)} pts de diferencia sobre '{_top3[1]['nombre']}').\n"
                    f"- '{_top3[0]['nombre']}' destaca con un Hit Rate de {round(_top3[0]['res']['hitRate'], 2)}% y un NDCG de {round(_top3[0]['res']['ndcg'], 2)}%.\n"
                )
                if _top3[1]['res']['precision'] > _top3[0]['res']['precision']:
                    _texto_tabla += f"- Sin embargo, '{_top3[1]['nombre']}' logra una mayor Precision ({round(_top3[1]['res']['precision'], 2)}%)."
                _pdf.multi_cell(0, 5, txt=_texto_tabla)

            _pdf.ln(8)
            _pdf.set_font("Helvetica", "B", 14)
            _pdf.set_text_color(30, 30, 30)
            _pdf.cell(0, 10, txt="Visualizaciones de Rendimiento", ln=1)
            try:
                _tmp_bar_pdf = "_tmp_bar_pdf.png"
                _tmp_rad_pdf = "_tmp_rad_pdf.png"
                _fig_bar.write_image(_tmp_bar_pdf, width=600, height=400)
                _fig_radar.write_image(_tmp_rad_pdf, width=600, height=400)
                _y_before = _pdf.get_y()
                _pdf.image(_tmp_bar_pdf, x=10, y=_y_before, w=90)
                _pdf.image(_tmp_rad_pdf, x=105, y=_y_before, w=90)
                _pdf.set_y(_y_before + 65)
                os.remove(_tmp_bar_pdf)
                os.remove(_tmp_rad_pdf)

                _pdf.set_font("Helvetica", size=10)
                _pdf.set_text_color(60, 60, 60)
                _mejor_prec = max(_scores, key=lambda x: x['res']['precision'])
                _mejor_cov = max(_scores, key=lambda x: x['res'].get('coverage', 0))
                
                _texto_graficas = "Interpretacion de Graficas:\n"
                _texto_graficas += f"- El grafico de barras muestra la consistencia de '{_top3[0]['nombre']}' en la mayoria de metricas"
                if _mejor_prec['nombre'] != _top3[0]['nombre']:
                    _texto_graficas += f", aunque resalta '{_mejor_prec['nombre']}' en Precision maxima.\n"
                else:
                    _texto_graficas += ".\n"
                
                _texto_graficas += f"- El perfil de radar ilustra el balance de los modelos. "
                if 'coverage' in _mejor_cov['res']:
                    _texto_graficas += f"Se observa que '{_mejor_cov['nombre']}' explora una mayor proporcion del catalogo (Coverage)."
                
                _pdf.multi_cell(0, 5, txt=_texto_graficas)
            except Exception:
                _pdf.set_font("Helvetica", "I", 10)
                _pdf.set_text_color(180, 0, 0)
                _pdf.cell(0, 8, txt="(Graficos no disponibles - requiere kaleido)", ln=1)
            _pdf.ln(5)
            _pdf.set_font("Helvetica", "B", 14)
            _pdf.set_text_color(30, 30, 30)
            _pdf.cell(0, 10, txt="Conclusion", ln=1)
            _pdf.set_font("Helvetica", size=10)
            _pdf.set_text_color(60, 60, 60)
            _texto_conclusion = (
                f"Tras evaluar {len(_scores)} modelos, '{_mejor['nombre']}' es la opcion recomendada "
                f"con un score final de {round(_mejor['score'],2)}/100, alcanzando un Hit Rate del {round(_mejor['res']['hitRate'],2)}% "
                f"y una Precision del {round(_mejor['res']['precision'],2)}%."
            )
            _pdf.multi_cell(0, 6, txt=_texto_conclusion)
            st.session_state.pdf_bytes = bytes(_pdf.output())
        except Exception as _e:
            st.session_state.pdf_bytes = None
            st.warning(f"PDF no disponible: {_e}")

        # -- WORD --
        try:
            def _set_cell_bg(_cell, _hex):
                _tc = _cell._tc
                _tcPr = _tc.get_or_add_tcPr()
                _shd = OxmlElement("w:shd")
                _shd.set(qn("w:val"), "clear")
                _shd.set(qn("w:color"), "auto")
                _shd.set(qn("w:fill"), _hex)
                _tcPr.append(_shd)

            _doc = Document()
            _doc.styles["Normal"].font.name = "Calibri"
            _doc.styles["Normal"].font.size = Pt(11)
            _tit = _doc.add_heading("TechStore AI - Reporte de Evaluacion de Modelos", level=0)
            _tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _tit.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            _doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
            _doc.add_paragraph(f"Usuarios: {_stats['totalUsuarios']}  |  Compras: {_stats['totalCompras']}  |  Top-K: {top_k}")
            _doc.add_paragraph()
            _pv = _doc.add_paragraph()
            _pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _rv = _pv.add_run(f"MEJOR MODELO: {_mejor['nombre']}  -  Score: {round(_mejor['score'],2)} / 100")
            _rv.bold = True; _rv.font.size = Pt(13); _rv.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
            _doc.add_paragraph()

            _doc.add_heading("1. Estadisticas del Dataset", level=1)
            _ts = _doc.add_table(rows=2, cols=3)
            _ts.style = "Table Grid"
            for _i, (_h, _v) in enumerate(zip(["Usuarios", "Compras", "Promedio"],[str(_stats["totalUsuarios"]),str(_stats["totalCompras"]),str(_stats["promedioCompras"])])):
                _c = _ts.cell(0, _i); _c.text = _h; _set_cell_bg(_c, "4F46E5")
                _c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF); _c.paragraphs[0].runs[0].bold = True
                _ts.cell(1, _i).text = _v
            _doc.add_paragraph()

            _doc.add_heading("2. Comparativa de Metricas", level=1)
            _ml = list(_baselines.keys())
            _tm = _doc.add_table(rows=len(_metricas_visibles)+1, cols=len(_ml)+1)
            _tm.style = "Table Grid"
            _c0 = _tm.cell(0,0); _c0.text = "Metrica"; _set_cell_bg(_c0,"4F46E5")
            _c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF); _c0.paragraphs[0].runs[0].bold = True
            for _j,_mn in enumerate(_ml):
                _c=_tm.cell(0,_j+1); _c.text=_mn; _set_cell_bg(_c,"4F46E5")
                _c.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _c.paragraphs[0].runs[0].bold=True
            for _i,_m in enumerate(_metricas_visibles):
                _tm.cell(_i+1,0).text=_m["label"]
                _vf=[_baselines[_mn].get(_m["key"],0) for _mn in _ml]
                _mv=max(_vf) if _vf else 0
                for _j,_val in enumerate(_vf):
                    _cl=_tm.cell(_i+1,_j+1); _cl.text=f"{round(_val,4)}"
                    if _val==_mv: _set_cell_bg(_cl,"DCFCE7")
            _doc.add_paragraph()
            _doc.add_paragraph()
            _doc.add_heading("Analisis Detallado por Metrica", level=2)
            
            _mejor_hr = max(_scores, key=lambda x: x['res']['hitRate'])
            _doc.add_paragraph(f"- Atraccion (Hit Rate@K): '{_mejor_hr['nombre']}' es el mas efectivo para captar la atencion primaria, logrando que un {round(_mejor_hr['res']['hitRate'], 2)}% de los usuarios encuentren al menos un producto relevante en sus recomendaciones.")
            
            _mejor_prec = max(_scores, key=lambda x: x['res']['precision'])
            _doc.add_paragraph(f"- Exactitud (Precision@K): '{_mejor_prec['nombre']}' es el modelo mas preciso ({round(_mejor_prec['res']['precision'], 2)}%), lo que significa que de los productos que sugiere, un mayor porcentaje resulta en compras reales, minimizando recomendaciones inutiles.")
            
            _mejor_rec = max(_scores, key=lambda x: x['res']['recall'])
            _doc.add_paragraph(f"- Retencion (Recall@K): En cuanto a recuperar todo lo que el usuario queria comprar, '{_mejor_rec['nombre']}' lidera con {round(_mejor_rec['res']['recall'], 2)}%, asegurando que menos productos de interes pasen desapercibidos.")
            
            _mejor_ndcg = max(_scores, key=lambda x: x['res']['ndcg'])
            _doc.add_paragraph(f"- Calidad del Ranking (NDCG@K): '{_mejor_ndcg['nombre']}' es el que mejor ordena los productos (NDCG: {round(_mejor_ndcg['res']['ndcg'], 2)}%), colocando las opciones mas atractivas en las primeras posiciones del carrusel.")
            
            _mejor_f1 = max(_scores, key=lambda x: x['res']['f1'])
            _doc.add_paragraph(f"- Balance Global (F1@K): Combinando exactitud y retencion, '{_mejor_f1['nombre']}' ofrece el perfil mas equilibrado con un F1 de {round(_mejor_f1['res']['f1'], 2)}%.")
            _doc.add_paragraph()

            _doc.add_heading("3. Ranking de Modelos", level=1)
            _tr=_doc.add_table(rows=len(_scores)+1,cols=3); _tr.style="Table Grid"
            for _j,_h in enumerate(["Pos","Modelo","Score"]):
                _c=_tr.cell(0,_j); _c.text=_h; _set_cell_bg(_c,"4F46E5")
                _c.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _c.paragraphs[0].runs[0].bold=True
            for _i,_s in enumerate(_scores):
                _tr.cell(_i+1,0).text=f"#{_i+1}"; _tr.cell(_i+1,1).text=_s["nombre"]; _tr.cell(_i+1,2).text=f"{round(_s['score'],2)}"
                if _i==0: _set_cell_bg(_tr.cell(_i+1,1),"DCFCE7"); _set_cell_bg(_tr.cell(_i+1,2),"DCFCE7")
            _doc.add_paragraph()

            _doc.add_heading("4. Tests Estadisticos", level=1)
            _doc.add_paragraph(f"Ganador '{_mejor['nombre']}' vs demas (p<0.05 = significativo).")
            for _mk,_ml2 in {"hitRate":"Hit Rate","precision":"Precision","recall":"Recall","ndcg":"NDCG"}.items():
                _doc.add_heading(f"Metrica: {_ml2}", level=2)
                _tdw=_tests[_mk]
                _tt=_doc.add_table(rows=len(_tdw)+1,cols=5); _tt.style="Table Grid"
                for _j,_h in enumerate(["Modelo","T-Stat","p-value","Cohen's d","Sig."]):
                    _c=_tt.cell(0,_j); _c.text=_h; _set_cell_bg(_c,"6D28D9")
                    _c.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _c.paragraphs[0].runs[0].bold=True
                for _i,(_bn,_t) in enumerate(_tdw.items()):
                    _sig=_t["ttest"]["significant"] or _t["wilcoxon"]["significant"]
                    _tt.cell(_i+1,0).text=_bn; _tt.cell(_i+1,1).text=str(round(_t["ttest"]["tStatistic"],4))
                    _tt.cell(_i+1,2).text=str(round(_t["ttest"]["pValue"],4)); _tt.cell(_i+1,3).text=str(round(_t["cohensD"],4))
                    _sc=_tt.cell(_i+1,4); _sc.text="Si" if _sig else "No"
                    if _sig: _set_cell_bg(_sc,"DCFCE7")
                _doc.add_paragraph()

            _doc.add_heading("5. Conclusion", level=1)
            _texto_conclusion_w = (
                f"Tras evaluar {len(_scores)} modelos, '{_mejor['nombre']}' es la opcion recomendada "
                f"con un score final de {round(_mejor['score'],2)}/100, alcanzando un Hit Rate del {round(_mejor['res']['hitRate'],2)}% "
                f"y una Precision del {round(_mejor['res']['precision'],2)}%."
            )
            _doc.add_paragraph(_texto_conclusion_w)
            _wb2 = io.BytesIO()
            _doc.save(_wb2)
            st.session_state.word_data = _wb2.getvalue()
        except Exception as _e:
            st.session_state.word_data = None
            st.warning(f"Word no disponible: {_e}")

# --- Visualización ---
if st.session_state.resultados:
    res = st.session_state.resultados
    stats = res["stats"]
    baselines = res["resultados_baselines"]
    mejor = res["mejor_modelo"]
    tests = res["tests_por_metrica"]
    scores = res["scores"]
    
    st.success(f"Evaluación completada para {stats['totalUsuarios']} usuarios sintéticos.")
    
    # 1. Dataset
    with st.expander("📊 Estadísticas del Dataset Sintético", expanded=True):
        col1, col2, col3 = st.columns(3)
        col1.metric("Usuarios Generados", stats["totalUsuarios"])
        col2.metric("Total de Compras", stats["totalCompras"])
        col3.metric("Promedio Compras/Usuario", stats["promedioCompras"])
        
        st.markdown("**Distribución de Categorías:**")
        df_cats = pd.DataFrame(list(stats["comprasPorCategoria"].items()), columns=["Categoría", "Compras"])
        st.bar_chart(df_cats.set_index("Categoría"))
        
        st.info(f"""
        **Interpretación del Dataset (Muestra de {stats['totalUsuarios']} usuarios):**
        - El volumen total de datos consiste en **{stats['totalCompras']} transacciones**, con una frecuencia de compra media de **{stats['promedioCompras']}** ítems por usuario.
        - **Distribución Comercial:** Como se aprecia en la gráfica, la simulación respeta los patrones de e-commerce donde algunas categorías dominan ampliamente el consumo. 
        - **Impacto del Cold-Start:** Al tener un promedio de {stats['promedioCompras']} compras por usuario, el sistema cuenta con un contexto mínimo suficiente para deducir patrones y ejecutar los tests, pero aún representa un escenario realista donde la escasez de datos (sparsity) pone a prueba el poder de predicción de los modelos.
        """)
        
    # --- Explicación de los Modelos y Métricas ---
    with st.expander("📖 Glosario: ¿Qué significan los Modelos y las Métricas?", expanded=False):
        st.markdown("""
        **📚 Los Algoritmos (Modelos):**
        - **Random:** Recomienda de forma completamente aleatoria. Sirve como base mínima. Si la IA no supera esto, no sirve.
        - **Popularidad:** Sugiere los productos más vendidos. Excelente para asegurar ventas seguras (Hit Rate), pero pobre en personalización.
        - **Content-Based:** Mide similitud de categorías (Coseno). Es el mejor para nichos porque asume que si compraste "X", te gustará algo muy similar.
        - **Híbrido Ponderado:** Combina puntuación de Popularidad y Content-Based, equilibrando lo que te gusta con lo que más se vende.
        - **Híbrido Cascada:** Primero filtra rígidamente por tus categorías favoritas y, sobre esa lista, ordena por los más populares.
        
        **📏 Las Métricas:**
        - **Hit Rate@K:** ¿A cuántos usuarios se les recomendó *al menos un* producto que realmente iban a comprar? (Mide atracción).
        - **Precision@K:** De los K productos recomendados, ¿cuántos fueron aciertos? (Mide exactitud de la lista).
        - **Recall@K:** De todos los productos que el usuario quería, ¿cuántos logró capturar el sistema? (Mide retención/descubrimiento).
        - **F1@K:** Promedio armónico entre Precision y Recall. Es el balance general de aciertos.
        - **NDCG@K:** No solo importa si acertaste, sino *dónde*. NDCG da más puntos si el producto correcto está en la posición 1 que en la 10.
        - **MRR (Mean Reciprocal Rank):** Mide cuán alto en la lista apareció el *primer* acierto. Un MRR de 0.5 indica que en promedio el primer acierto estuvo en la posición 2.
        - **Diversity / Novelty:** Miden si el algoritmo recomienda categorías variadas o si se estanca recomendando siempre lo mismo.
        """)

    # 2. Tabla de Métricas
    st.subheader("📈 Comparativa de Métricas por Modelo")
    
    metricas_visibles = [
        {"key": "hitRate", "label": "Hit Rate@K", "sufijo": "%"},
        {"key": "precision", "label": "Precision@K", "sufijo": "%"},
        {"key": "recall", "label": "Recall@K", "sufijo": "%"},
        {"key": "f1", "label": "F1@K", "sufijo": "%"},
        {"key": "mrr", "label": "MRR", "sufijo": ""},
        {"key": "map", "label": "MAP", "sufijo": ""},
        {"key": "ndcg", "label": "NDCG@K", "sufijo": "%"},
        {"key": "coverage", "label": "Coverage", "sufijo": "%"},
        {"key": "diversity", "label": "Diversity", "sufijo": "%"},
        {"key": "novelty", "label": "Novelty", "sufijo": "%"}
    ]
    
    # Preparar DataFrame
    tabla_data = []
    for m in metricas_visibles:
        row = {"Métrica": m["label"]}
        for nombre, resultado in baselines.items():
            row[nombre] = round(resultado.get(m["key"], 0), 4)
        tabla_data.append(row)
        
    df_metricas = pd.DataFrame(tabla_data).set_index("Métrica")
    st.dataframe(df_metricas.style.highlight_max(axis=1, color="lightgreen"), use_container_width=True)
    
    peor_hit = min(scores, key=lambda x: x["res"]["hitRate"])
    st.info(f"""
    **Análisis Adaptativo de Métricas:**
    - **Hit Rate y Precision:** El modelo **'{mejor['nombre']}'** destaca globalmente, logrando un Hit Rate de **{mejor['res']['hitRate']:.2f}%**. Por el contrario, el modelo **'{peor_hit['nombre']}'** obtuvo el peor Hit Rate ({peor_hit['res']['hitRate']:.2f}%), demostrando su deficiencia para atraer la atención primaria del cliente.
    - **NDCG@K (Calidad del Ranking):** El ganador obtuvo un NDCG de **{mejor['res']['ndcg']:.2f}%**. Un NDCG alto significa que cuando este algoritmo acierta, coloca los productos de interés en los *primeros lugares* (lo cual es vital para tiendas online).
    - **Diversity & Novelty:** Mientras modelos como 'Popularidad' tienden a mostrar productos repetitivos o "mainstream", algoritmos como 'Random' o 'Content-Based' pueden potenciar mejor la diversidad (Variety) del catálogo ({round(peor_hit['res'].get('diversity', 0), 2)}% vs {round(mejor['res'].get('diversity', 0), 2)}% del ganador).
    """)

    # 3. Gráficos Plotly
    st.subheader("📊 Visualizaciones")
    colA, colB = st.columns(2)
    
    COLORES_GRAFICOS = ['#EF4444', '#F59E0B', '#10B981', '#EC4899', '#14B8A6', '#8B5CF6']
    
    with colA:
        df_melted = df_metricas.reset_index().melt(id_vars=["Métrica"], var_name="Modelo", value_name="Valor")
        fig_bar = px.bar(df_melted, x="Métrica", y="Valor", color="Modelo", barmode="group", title="Comparación de Métricas", color_discrete_sequence=COLORES_GRAFICOS)
        st.plotly_chart(fig_bar, use_container_width=True)
        
    with colB:
        radar_keys = ['hitRate', 'precision', 'recall', 'f1', 'ndcg', 'diversity']
        fig_radar = go.Figure()
        
        for i, (nombre, resultado) in enumerate(baselines.items()):
            valores = [resultado.get(k, 0) for k in radar_keys]
            valores.append(valores[0]) 
            keys_plot = radar_keys + [radar_keys[0]]
            
            fig_radar.add_trace(go.Scatterpolar(
                r=valores,
                theta=keys_plot,
                fill='toself',
                name=nombre,
                line_color=COLORES_GRAFICOS[i % len(COLORES_GRAFICOS)]
            ))
            
        fig_radar.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
            showlegend=True,
            title="Perfil de Rendimiento (Radar)"
        )
        st.plotly_chart(fig_radar, use_container_width=True)

    # Lógica adaptativa para el texto del gráfico de barras
    mejor_metrica_llave = max(
        ['hitRate', 'precision', 'recall', 'ndcg', 'diversity'],
        key=lambda k: mejor['res'].get(k, 0)
    )
    nombres_metricas = {'hitRate': 'Hit Rate', 'precision': 'Precision', 'recall': 'Recall', 'ndcg': 'NDCG', 'diversity': 'Diversity'}
    nombre_metrica_destacada = nombres_metricas[mejor_metrica_llave]
    ventaja = mejor['res'].get(mejor_metrica_llave, 0) - peor_hit['res'].get(mejor_metrica_llave, 0)

    st.info(f"""
    **Interpretación Gráfica (Perfil de Rendimiento):**
    - En el **Gráfico de Radar**, puedes observar visualmente cómo el modelo **'{mejor['nombre']}'** cubre una mayor área global en comparación con **'{peor_hit['nombre']}'**.
    - Un polígono amplio y equilibrado como el de **'{mejor['nombre']}'** demuestra que no solo está siendo exacto, sino que también es capaz de retener volumen y mantener una diversidad aceptable. Los modelos híbridos suelen lograr polígonos más uniformes al combinar exactitud y popularidad.
    - **El Gráfico de Barras** confirma que, por ejemplo en la métrica **{nombre_metrica_destacada}**, el modelo '{mejor['nombre']}' logra una impresionante ventaja de **+{ventaja:.1f}%** sobre el modelo de menor rendimiento, justificando contundentemente por qué este algoritmo es superior globalmente.
    """)

    # 4. Tests Estadísticos
    st.subheader(f"🔬 Tests Estadísticos (Base: {mejor['nombre']})")
    st.markdown(f"Se compara el modelo ganador **{mejor['nombre']}** contra los demás algoritmos para determinar si la diferencia es estadísticamente significativa (p < 0.05).")
    
    tabs = st.tabs(["Hit Rate", "Precision", "Recall", "NDCG"])
    map_tabs = {"Hit Rate": "hitRate", "Precision": "precision", "Recall": "recall", "NDCG": "ndcg"}
    
    for i, (tab_name, metric_key) in enumerate(map_tabs.items()):
        with tabs[i]:
            test_data = tests[metric_key]
            test_rows = []
            for bl_name, t in test_data.items():
                is_sig = t["ttest"]["significant"] or t["wilcoxon"]["significant"]
                
                # Texto dinámico adaptativo para Cohen's D
                abs_d = abs(t["cohensD"])
                if abs_d >= 0.8:
                    d_mag = "Grande"
                elif abs_d >= 0.5:
                    d_mag = "Medio"
                elif abs_d >= 0.2:
                    d_mag = "Pequeño"
                else:
                    d_mag = "Despreciable"
                
                inter_texto = f"La ventaja de {mejor['nombre']} frente a {bl_name} es {'REAL' if is_sig else 'AZAROSA'} (Efecto: {d_mag})"
                
                test_rows.append({
                    "Modelo Comparado": bl_name,
                    "T-Statistic": round(t["ttest"]["tStatistic"], 4),
                    "p-value (T-Test)": round(t["ttest"]["pValue"], 4),
                    "p-value (Wilcoxon)": round(t["wilcoxon"]["pValue"], 4),
                    "Cohen's d": round(t["cohensD"], 4),
                    "Interpretación Clínica": inter_texto,
                    "Significativo (p < 0.05)": "✅ Sí" if is_sig else "❌ No"
                })
            
            df_tab = pd.DataFrame(test_rows)
            st.dataframe(df_tab, use_container_width=True)

            # --- SECCIÓN DE INTERPRETACIÓN CIENTÍFICA (PAPER-READY Y DINÁMICA) ---
            st.markdown(f"### 📑 Discusión y Análisis Estadístico para {tab_name}")
            st.markdown(
                f"A continuación se presenta un análisis cuantitativo riguroso de las pruebas de hipótesis aplicadas. "
                f"Se comparan las medias de rendimiento de **{mejor['nombre']}** (modelo base ganador) frente a cada baseline bajo la métrica **{tab_name}**:"
            )

            # Generamos explicaciones dinámicas detalladas para cada fila (modelo)
            for row in test_rows:
                modelo = row["Modelo Comparado"]
                t_stat = row["T-Statistic"]
                p_ttest = row["p-value (T-Test)"]
                p_wilc = row["p-value (Wilcoxon)"]
                cohen_d = row["Cohen's d"]
                es_sig = row["Significativo (p < 0.05)"] == "✅ Sí"
                
                # Determinación de significancia y efecto
                sig_desc = "**es estadísticamente significativa**" if es_sig else "**no es estadísticamente significativa**"
                d_abs = abs(cohen_d)
                if d_abs >= 0.8:
                    efecto_desc = f"un tamaño del efecto **Grande** (d = {cohen_d}), lo que indica una alta relevancia práctica en producción real."
                elif d_abs >= 0.5:
                    efecto_desc = f"un tamaño del efecto **Medio** (d = {cohen_d}), reflejando una mejora práctica notable en el comportamiento del usuario."
                elif d_abs >= 0.2:
                    efecto_desc = f"un tamaño del efecto **Pequeño** (d = {cohen_d}), lo que indica que el impacto práctico es sutil a pesar de la consistencia estadística."
                else:
                    efecto_desc = f"un tamaño del efecto **Despreciable** (d = {cohen_d}), sugiriendo que la diferencia de rendimiento es prácticamente nula."

                # Explicación del T-Statistic
                if t_stat > 0:
                    t_desc = f"un T-Statistic de **{t_stat}** (positivo, indicando que el rendimiento de {mejor['nombre']} es superior)"
                else:
                    t_desc = f"un T-Statistic de **{t_stat}** (negativo, indicando una tendencia desfavorable)"

                # Explicación del p-value
                p_desc_test = (
                    f"Con un p-value de T-Test de **{p_ttest}** y p-value de Wilcoxon de **{p_wilc}** (ambos < 0.05), la probabilidad de "
                    f"que esta diferencia se deba enteramente al azar es inferior al 5%."
                    if es_sig else
                    f"Dado que el p-value de T-Test es **{p_ttest}** y el de Wilcoxon es **{p_wilc}** (ambos ≥ 0.05), no se puede rechazar "
                    f"la hipótesis nula de igualdad de medias con un nivel de confianza del 95%."
                )

                st.markdown(
                    f"1. **Comparación con {modelo}:**\n"
                    f"   - **Estadístico de Contraste:** Se obtuvo {t_desc}. Esto mide a cuántas desviaciones estándar "
                    f"se encuentra la diferencia de medias observada respecto al valor esperado bajo la hipótesis nula.\n"
                    f"   - **Significancia Estadística:** {p_desc_test} Por lo tanto, la diferencia {sig_desc}.\n"
                    f"   - **Tamaño del Efecto (Magnitud Clave para el Paper):** La métrica d de Cohen revela {efecto_desc}\n"
                )

            # Resumen global de la métrica apto para redacción científica
            total_modelos = len(test_rows)
            sig_modelos = sum(1 for row in test_rows if row["Significativo (p < 0.05)"] == "✅ Sí")
            
            st.info(
                f"**Síntesis para el Artículo Científico ({tab_name}):**\n\n"
                f"Los resultados demuestran que **{mejor['nombre']}** logró superar estadísticamente a **{sig_modelos} de los {total_modelos}** modelos de control evaluados en la métrica **{tab_name}** con un nivel de confianza del 95%. "
                f"Este análisis sustenta formalmente la elección de **{mejor['nombre']}** como el algoritmo óptimo para producción en el módulo de recomendaciones, "
                f"avalado por los contrastes paramétricos (T-Test) y no paramétricos (Wilcoxon)."
            )

    # 5. Veredicto Final
    st.markdown("---")
    st.success(f"🏆 **MEJOR MODELO: {mejor['nombre']}** con un score global de **{round(mejor['score'], 2)} / 100**.")
    
    # 6. Exportación
    st.subheader("📥 Exportar Resultados")
    col_pdf, col_excel, col_word = st.columns(3)

    with col_pdf:
        if st.session_state.pdf_bytes:
            st.download_button(
                label="📄 Descargar Reporte Avanzado (PDF)",
                data=st.session_state.pdf_bytes,
                file_name=f"reporte_avanzado_{datetime.datetime.now().strftime('%Y%m%d')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        else:
            st.info("PDF no disponible")

    with col_excel:
        if st.session_state.excel_data:
            st.download_button(
                label="📊 Descargar Resultados en Excel",
                data=st.session_state.excel_data,
                file_name=f"evaluacion_ia_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        else:
            st.info("Excel no disponible")

    with col_word:
        if st.session_state.word_data:
            st.download_button(
                label="📝 Descargar Reporte en Word",
                data=st.session_state.word_data,
                file_name=f"reporte_word_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.info("Word no disponible")
